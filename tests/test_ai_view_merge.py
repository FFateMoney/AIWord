"""Tests for ai_view, ai_merge, and _inherit_style_rPr."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from word_ast import parse_docx, to_ai_view, merge_ai_edits


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _make_minimal_ast(para_fmt=None, run_overrides=None):
    """Build a minimal AST dict for merge/view tests."""
    content = [{"type": "Text", "text": "Hello"}]
    if run_overrides:
        content[0]["overrides"] = run_overrides
    block = {"id": "p0", "type": "Paragraph", "content": content}
    if para_fmt:
        block["paragraph_format"] = para_fmt
    return {
        "schema_version": "1.0",
        "document": {"body": [block]},
    }


# ---------------------------------------------------------------------------
# to_ai_view tests
# ---------------------------------------------------------------------------

def test_to_ai_view_strips_raw_pPr():
    """to_ai_view must remove _raw_pPr from paragraph_format."""
    raw_xml = f'<w:pPr xmlns:w="{_W_NS}"><w:jc w:val="center"/></w:pPr>'
    ast = _make_minimal_ast(para_fmt={"alignment": "center", "_raw_pPr": raw_xml})
    view = to_ai_view(ast)
    para = view["document"]["body"][0]
    assert "_raw_pPr" not in para.get("paragraph_format", {})
    assert para["paragraph_format"]["alignment"] == "center"


def test_to_ai_view_strips_raw_rPr():
    """to_ai_view must remove _raw_rPr from run overrides."""
    raw_xml = f'<w:rPr xmlns:w="{_W_NS}"><w:b/></w:rPr>'
    ast = _make_minimal_ast(run_overrides={"bold": True, "_raw_rPr": raw_xml})
    view = to_ai_view(ast)
    run = view["document"]["body"][0]["content"][0]
    assert "_raw_rPr" not in run.get("overrides", {})
    assert run["overrides"]["bold"] is True


def test_to_ai_view_does_not_mutate_original():
    """to_ai_view must not modify the original AST."""
    raw_xml = f'<w:rPr xmlns:w="{_W_NS}"><w:b/></w:rPr>'
    ast = _make_minimal_ast(run_overrides={"bold": True, "_raw_rPr": raw_xml})
    to_ai_view(ast)
    run = ast["document"]["body"][0]["content"][0]
    assert "_raw_rPr" in run.get("overrides", {})


def test_to_ai_view_preserves_non_raw_fields():
    """to_ai_view must keep all semantic (non _raw_) fields intact."""
    ast = _make_minimal_ast(
        para_fmt={"alignment": "right", "indent_left": 720},
        run_overrides={"bold": True, "size": 24, "color": "#FF0000"},
    )
    view = to_ai_view(ast)
    para = view["document"]["body"][0]
    assert para["paragraph_format"] == {"alignment": "right", "indent_left": 720}
    assert para["content"][0]["overrides"] == {"bold": True, "size": 24, "color": "#FF0000"}


# ---------------------------------------------------------------------------
# merge_ai_edits tests
# ---------------------------------------------------------------------------

def test_merge_preserves_raw_pPr_when_para_fmt_unchanged():
    """merge_ai_edits must keep _raw_pPr when the AI does not change paragraph_format."""
    raw_xml = f'<w:pPr xmlns:w="{_W_NS}"><w:jc w:val="center"/></w:pPr>'
    original = _make_minimal_ast(para_fmt={"alignment": "center", "_raw_pPr": raw_xml})
    ai_view = _make_minimal_ast(para_fmt={"alignment": "center"})
    result = merge_ai_edits(original, ai_view)
    fmt = result["document"]["body"][0]["paragraph_format"]
    assert "_raw_pPr" in fmt
    assert fmt["_raw_pPr"] == raw_xml


def test_merge_updates_alignment_in_raw_pPr():
    """merge_ai_edits must update <w:jc> in _raw_pPr when alignment changes."""
    raw_xml = f'<w:pPr xmlns:w="{_W_NS}"><w:jc w:val="center"/></w:pPr>'
    original = _make_minimal_ast(para_fmt={"alignment": "center", "_raw_pPr": raw_xml})
    ai_view = _make_minimal_ast(para_fmt={"alignment": "right"})
    result = merge_ai_edits(original, ai_view)
    fmt = result["document"]["body"][0]["paragraph_format"]
    assert fmt["alignment"] == "right"
    assert "_raw_pPr" in fmt
    assert 'w:val="right"' in fmt["_raw_pPr"]


def test_merge_preserves_raw_rPr_when_run_overrides_unchanged():
    """merge_ai_edits must keep _raw_rPr when the AI does not change run overrides."""
    raw_xml = f'<w:rPr xmlns:w="{_W_NS}"><w:b/></w:rPr>'
    original = _make_minimal_ast(run_overrides={"bold": True, "_raw_rPr": raw_xml})
    ai_view = _make_minimal_ast(run_overrides={"bold": True})
    result = merge_ai_edits(original, ai_view)
    ov = result["document"]["body"][0]["content"][0]["overrides"]
    assert "_raw_rPr" in ov
    assert ov["_raw_rPr"] == raw_xml


def test_merge_updates_bold_in_raw_rPr():
    """merge_ai_edits must add <w:b/> to _raw_rPr when bold is set to True."""
    raw_xml = f'<w:rPr xmlns:w="{_W_NS}"></w:rPr>'
    original = _make_minimal_ast(run_overrides={"_raw_rPr": raw_xml})
    ai_view = _make_minimal_ast(run_overrides={"bold": True})
    result = merge_ai_edits(original, ai_view)
    ov = result["document"]["body"][0]["content"][0]["overrides"]
    assert ov["bold"] is True
    assert "_raw_rPr" in ov
    assert "<w:b" in ov["_raw_rPr"]


def test_merge_removes_bold_from_raw_rPr():
    """merge_ai_edits must remove <w:b/> from _raw_rPr when bold is set to False."""
    raw_xml = f'<w:rPr xmlns:w="{_W_NS}"><w:b/></w:rPr>'
    original = _make_minimal_ast(run_overrides={"bold": True, "_raw_rPr": raw_xml})
    ai_view = _make_minimal_ast(run_overrides={"bold": False})
    result = merge_ai_edits(original, ai_view)
    ov = result["document"]["body"][0]["content"][0]["overrides"]
    assert ov["bold"] is False
    assert "<w:b" not in ov.get("_raw_rPr", "")


def test_merge_updates_font_in_raw_rPr():
    """merge_ai_edits must update <w:rFonts> when font_ascii changes."""
    raw_xml = (
        f'<w:rPr xmlns:w="{_W_NS}">'
        '<w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
        "</w:rPr>"
    )
    original = _make_minimal_ast(run_overrides={"font_ascii": "Arial", "_raw_rPr": raw_xml})
    ai_view = _make_minimal_ast(run_overrides={"font_ascii": "SimSun"})
    result = merge_ai_edits(original, ai_view)
    ov = result["document"]["body"][0]["content"][0]["overrides"]
    assert ov["font_ascii"] == "SimSun"
    assert "SimSun" in ov["_raw_rPr"]


def test_merge_updates_size_in_raw_rPr():
    """merge_ai_edits must update <w:sz> and <w:szCs> when size changes."""
    raw_xml = (
        f'<w:rPr xmlns:w="{_W_NS}">'
        '<w:sz w:val="24"/><w:szCs w:val="24"/>'
        "</w:rPr>"
    )
    original = _make_minimal_ast(run_overrides={"size": 24, "_raw_rPr": raw_xml})
    ai_view = _make_minimal_ast(run_overrides={"size": 28})
    result = merge_ai_edits(original, ai_view)
    ov = result["document"]["body"][0]["content"][0]["overrides"]
    assert ov["size"] == 28
    assert 'w:val="28"' in ov["_raw_rPr"]


def test_merge_drops_raw_on_xml_parse_failure():
    """merge_ai_edits must drop _raw_pPr when the XML is unparseable."""
    original = _make_minimal_ast(
        para_fmt={"alignment": "center", "_raw_pPr": "NOT VALID XML"}
    )
    ai_view = _make_minimal_ast(para_fmt={"alignment": "right"})
    result = merge_ai_edits(original, ai_view)
    fmt = result["document"]["body"][0]["paragraph_format"]
    assert "_raw_pPr" not in fmt
    assert fmt["alignment"] == "right"


def test_merge_updates_text():
    """merge_ai_edits must update run text when the AI changes it."""
    original = _make_minimal_ast()
    ai_view = _make_minimal_ast()
    ai_view["document"]["body"][0]["content"][0]["text"] = "World"
    result = merge_ai_edits(original, ai_view)
    assert result["document"]["body"][0]["content"][0]["text"] == "World"


def test_merge_unmatched_block_preserved():
    """merge_ai_edits must leave blocks not present in ai_ast unchanged."""
    raw_xml = f'<w:pPr xmlns:w="{_W_NS}"><w:jc w:val="center"/></w:pPr>'
    original = {
        "schema_version": "1.0",
        "document": {
            "body": [
                {
                    "id": "p0",
                    "type": "Paragraph",
                    "paragraph_format": {"alignment": "center", "_raw_pPr": raw_xml},
                    "content": [],
                }
            ]
        },
    }
    ai_view = {"document": {"body": []}}  # AI returned empty body
    result = merge_ai_edits(original, ai_view)
    fmt = result["document"]["body"][0]["paragraph_format"]
    assert "_raw_pPr" in fmt


# ---------------------------------------------------------------------------
# _inherit_style_rPr integration tests (via parse_docx)
# ---------------------------------------------------------------------------

def test_inherit_style_rPr_captures_style_font(tmp_path: Path):
    """Runs without explicit font must capture style-inherited font in _raw_rPr."""
    path = tmp_path / "inherit_font.docx"
    doc = Document()

    style = doc.styles["Heading 1"]
    style.font.name = "SimSun"
    style.font.size = Pt(14)

    p = doc.add_paragraph("Test text")
    p.style = style
    # The run carries no explicit font — font comes from the style
    doc.save(path)

    ast = parse_docx(path)
    para_block = next(
        b for b in ast["document"]["body"] if b.get("type") == "Paragraph"
    )
    run = para_block["content"][0]
    raw_rPr = run.get("overrides", {}).get("_raw_rPr", "")
    # _raw_rPr must contain the font information inherited from "Heading 1"
    assert "SimSun" in raw_rPr


def test_inherit_pPr_no_pPr_element(tmp_path: Path):
    """Paragraphs without an explicit <w:pPr> still capture inherited pPr."""
    path = tmp_path / "no_pPr.docx"
    doc = Document()

    # Use a style that defines alignment so the paragraph inherits it
    style = doc.styles["Heading 1"]
    # python-docx Heading 1 typically has its own pPr; use a paragraph style
    # that we customise with alignment so we can verify inheritance
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("No pPr")
    p.style = style
    doc.save(path)

    ast = parse_docx(path)
    para_block = next(
        b for b in ast["document"]["body"] if b.get("type") == "Paragraph"
    )
    # Either _raw_pPr contains the inherited jc or structural alignment is set
    para_fmt = para_block.get("paragraph_format", {})
    raw_pPr = para_fmt.get("_raw_pPr", "")
    alignment = para_fmt.get("alignment")
    assert "center" in raw_pPr or alignment == "center"


def test_ai_view_merge_roundtrip_preserves_text(tmp_path: Path):
    """Full pipeline: parse → ai_view → merge (no changes) → render → same text."""
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    p = doc.add_paragraph("Hello AI")
    p.add_run(" World").bold = True
    doc.save(src)

    from word_ast import render_ast

    ast = parse_docx(src)
    ai_view = to_ai_view(ast)
    merged = merge_ai_edits(ast, ai_view)
    render_ast(merged, out)

    rebuilt = Document(out)
    assert rebuilt.paragraphs[0].text == "Hello AI World"
