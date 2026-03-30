import base64
import io
from pathlib import Path

from docx import Document
from docx.shared import Inches

from word_ast import parse_docx, to_content_view


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
    "z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
)


def test_to_content_view_strips_formatting_fields():
    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {"page": {"size": "A4"}},
            "styles": {
                "Heading1": {"style_id": "Heading1", "name": "Heading 1", "type": "paragraph"}
            },
            "body": [
                {
                    "id": "p0",
                    "type": "Paragraph",
                    "style": "Heading1",
                    "paragraph_format": {"alignment": "center", "_raw_pPr": "<w:pPr />"},
                    "default_run": {"size": 24},
                    "content": [
                        {
                            "type": "Text",
                            "text": "Title",
                            "overrides": {"bold": True, "_raw_rPr": "<w:rPr />"},
                        }
                    ],
                }
            ],
            "passthrough": {"header_xml": "<w:hdr />"},
        },
    }

    view = to_content_view(ast)

    assert "meta" not in view["document"]
    assert "styles" not in view["document"]
    heading = view["document"]["body"][0]
    assert heading == {
        "id": "p0",
        "type": "Heading",
        "level": 1,
        "text": "Title",
        "content": [{"type": "Text", "text": "Title"}],
    }


def test_to_content_view_preserves_table_structure():
    ast = {
        "schema_version": "1.0",
        "document": {
            "styles": {},
            "body": [
                {
                    "id": "t0",
                    "type": "Table",
                    "rows": [
                        {
                            "cells": [
                                {
                                    "id": "t0.r0c0",
                                    "col_span": 2,
                                    "row_span": 1,
                                    "content": [
                                        {
                                            "id": "t0.r0c0.p0",
                                            "type": "Paragraph",
                                            "content": [{"type": "Text", "text": "Cell A"}],
                                        }
                                    ],
                                    "_raw_tcPr": "<w:tcPr />",
                                }
                            ],
                            "_raw_trPr": "<w:trPr />",
                        }
                    ],
                    "_raw_tblPr": "<w:tblPr />",
                }
            ],
        },
    }

    view = to_content_view(ast)
    table = view["document"]["body"][0]

    assert table["type"] == "Table"
    assert table["rows"][0]["cells"][0]["text"] == "Cell A"
    assert table["rows"][0]["cells"][0]["content"][0]["type"] == "Paragraph"
    assert "_raw_tblPr" not in table


def test_to_content_view_keeps_inline_image_binary(tmp_path: Path):
    path = tmp_path / "image.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(_PNG_1X1), width=Inches(1), height=Inches(1))
    doc.save(path)

    ast = parse_docx(path)
    view = to_content_view(ast)

    para = view["document"]["body"][0]
    image = next(item for item in para["content"] if item["type"] == "InlineImage")
    assert image["data"]
    assert image["content_type"] == "image/png"
    assert image["width"] > 0
    assert image["height"] > 0
