from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from word_ast.parser.document_parser import parse_docx


def test_parse_character_format_bold(tmp_path: Path):
    path = tmp_path / "bold.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("normal ")
    run = p.add_run("bold")
    run.bold = True
    doc.save(path)

    ast = parse_docx(path)
    runs = ast["document"]["body"][0]["content"]
    bold_run = next(r for r in runs if r.get("overrides", {}).get("bold"))
    assert bold_run["overrides"]["bold"] is True


def test_parse_styles_type_is_normalized(tmp_path: Path):
    path = tmp_path / "styles.docx"
    doc = Document()
    doc.add_paragraph("hello")
    doc.save(path)

    ast = parse_docx(path)
    styles = ast["document"]["styles"]
    assert styles["Normal"]["type"] == "paragraph"


def test_parse_merged_table_spans(tmp_path: Path):
    path = tmp_path / "merged.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    for row_idx in range(3):
        for col_idx in range(3):
            table.cell(row_idx, col_idx).text = f"{row_idx},{col_idx}"

    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 2).merge(table.cell(2, 2))
    doc.save(path)

    ast = parse_docx(path)
    table_ast = next(block for block in ast["document"]["body"] if block["type"] == "Table")

    merged_h = table_ast["rows"][0]["cells"][0]
    assert merged_h["col_span"] == 2
    assert merged_h["row_span"] == 1

    merged_v = table_ast["rows"][1]["cells"][2]
    assert merged_v["col_span"] == 1
    assert merged_v["row_span"] == 2

    assert len(table_ast["rows"][2]["cells"]) == 2


def test_parse_merges_consecutive_runs_with_same_style(tmp_path: Path):
    """Consecutive runs with identical formatting should be merged into one Text node."""
    path = tmp_path / "fragmented.docx"
    doc = Document()
    p = doc.add_paragraph()
    # Three runs with the same formatting — simulates DOCX fragmentation
    for text in ["核心网", "来管理5G基站。", "核心网"]:
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        r.font.size = Pt(12)
        r.font.name = "宋体"
    doc.save(path)

    ast = parse_docx(path)
    content = ast["document"]["body"][0]["content"]
    assert len(content) == 1
    assert content[0]["text"] == "核心网来管理5G基站。核心网"


def test_parse_does_not_merge_runs_with_different_styles(tmp_path: Path):
    """Runs with different formatting must stay separate after parsing."""
    path = tmp_path / "diff.docx"
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("red ")
    r1.font.color.rgb = RGBColor(255, 0, 0)
    r2 = p.add_run("blue")
    r2.font.color.rgb = RGBColor(0, 0, 255)
    doc.save(path)

    ast = parse_docx(path)
    content = ast["document"]["body"][0]["content"]
    assert len(content) == 2
