import base64
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from lxml import etree

from word_ast import parse_docx, render_ast

# Minimal 1×1 transparent PNG used in image round-trip tests
_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8"
    "z8BQDwADhQGAWjR9awAAAABJRU5ErkJggg=="
)


def test_roundtrip_text_and_table(tmp_path: Path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Hello")
    run.bold = True
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(src)

    ast = parse_docx(src)
    assert ast["schema_version"] == "1.0"
    assert ast["document"]["body"][0]["type"] == "Paragraph"
    assert ast["document"]["body"][1]["type"] == "Table"

    render_ast(ast, out)

    rebuilt = Document(out)
    assert rebuilt.paragraphs[0].text == "Hello"
    assert rebuilt.tables[0].cell(0, 0).text == "A"
    assert rebuilt.tables[0].cell(0, 1).text == "B"


def test_render_uses_style_name_fallback_for_style_id(tmp_path: Path):
    out = tmp_path / "out-heading.docx"

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {
                "1": {"style_id": "1", "name": "Heading 1", "type": "paragraph", "based_on": "Normal"},
                "a": {"style_id": "a", "name": "Normal", "type": "paragraph", "based_on": None},
            },
            "body": [
                {"id": "p0", "type": "Paragraph", "style": "1", "content": [{"type": "Text", "text": "Title"}]},
                {"id": "p1", "type": "Paragraph", "style": "a", "content": [{"type": "Text", "text": "Body"}]},
            ],
            "passthrough": {},
        },
    }

    render_ast(ast, out)

    rebuilt = Document(out)
    assert rebuilt.paragraphs[0].style.name == "Heading 1"
    assert rebuilt.paragraphs[1].style.name == "Normal"


def test_roundtrip_preserves_paragraph_style_font_defaults(tmp_path: Path):
    """Runs without explicit formatting inherit from the paragraph style.
    The paragraph style is applied correctly and the text content is preserved."""
    src = tmp_path / "styled.docx"
    out = tmp_path / "styled-out.docx"

    doc = Document()
    heading_style = doc.styles["Heading 1"]
    heading_style.font.name = "Arial"
    heading_style.font.size = Pt(24)
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.font.bold = True

    heading = doc.add_paragraph("Styled title")
    heading.style = heading_style
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    para = rebuilt.paragraphs[0]
    assert para.text == "Styled title"
    assert para.style.name == "Heading 1"
    # Runs without explicit formatting should NOT have explicit overrides;
    # they inherit from the paragraph style, avoiding fake-bold artifacts.
    run = para.runs[0]
    assert run.font.name is None


def _get_east_asia_font(run) -> str | None:
    """Helper to read the East Asian font name from a run's XML."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return None
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        return None
    return rFonts.get(qn('w:eastAsia'))


def test_roundtrip_preserves_east_asian_font(tmp_path: Path):
    """East Asian font (w:eastAsia) must survive a parse → render round-trip."""
    src = tmp_path / "ea.docx"
    out = tmp_path / "ea-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("你好世界")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    # Set East Asian font via XML
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '宋体')
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    r = rebuilt.paragraphs[0].runs[0]
    assert r.font.name == "Calibri"
    assert _get_east_asia_font(r) == "宋体"
    assert r.font.size.pt == 14


def test_roundtrip_multi_run_different_fonts(tmp_path: Path):
    """Multiple runs with different fonts/colors/sizes must not bleed into each other."""
    src = tmp_path / "multi.docx"
    out = tmp_path / "multi-out.docx"

    doc = Document()
    p = doc.add_paragraph()

    r1 = p.add_run("Red ")
    r1.font.color.rgb = RGBColor(255, 0, 0)
    r1.font.size = Pt(14)
    r1.font.name = "Arial"

    r2 = p.add_run("Blue")
    r2.font.color.rgb = RGBColor(0, 0, 255)
    r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"

    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    runs = rebuilt.paragraphs[0].runs
    assert len(runs) == 2

    assert runs[0].font.name == "Arial"
    assert runs[0].font.size.pt == 14
    assert str(runs[0].font.color.rgb) == "FF0000"

    assert runs[1].font.name == "Times New Roman"
    assert runs[1].font.size.pt == 10
    assert str(runs[1].font.color.rgb) == "0000FF"


def test_roundtrip_style_defaults_with_run_overrides(tmp_path: Path):
    """When a style sets defaults and a run overrides only some properties,
    only the explicitly overridden properties should be present on the run.
    Non-overridden properties come from the paragraph style, avoiding
    fake-bold artifacts and formatting corruption."""
    src = tmp_path / "override.docx"
    out = tmp_path / "override-out.docx"

    doc = Document()
    heading = doc.styles["Heading 1"]
    heading.font.name = "Arial"
    heading.font.size = Pt(24)
    heading.font.color.rgb = RGBColor(0, 0, 128)
    heading.font.bold = True
    # Set East Asian font on style
    rPr = heading.font._element.rPr
    rFonts = rPr.find(qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.style = heading

    # Run 1: no overrides (inherit all from style)
    p.add_run("Title ")

    # Run 2: override only color
    r2 = p.add_run("RED")
    r2.font.color.rgb = RGBColor(255, 0, 0)

    # Run 3: override only ASCII font (east-asian font must come from style default)
    r3 = p.add_run(" Serif")
    r3.font.name = "Times New Roman"

    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    para = rebuilt.paragraphs[0]
    assert para.style.name == "Heading 1"
    runs = para.runs

    # Run 1: no explicit overrides — inherits everything from the style
    assert runs[0].font.name is None
    assert runs[0].bold is None

    # Run 2: only color overridden; other properties inherit from style
    assert str(runs[1].font.color.rgb) == "FF0000"

    # Run 3: only ASCII font overridden
    assert runs[2].font.name == "Times New Roman"


def _assert_no_heading_colors_in_styles(styles_element):
    """Assert no heading style in *styles_element* has a ``<w:color>``."""
    for style_el in styles_element.iterchildren(qn("w:style")):
        name_el = style_el.find(qn("w:name"))
        if name_el is None:
            continue
        name_val = name_el.get(qn("w:val"), "")
        if "heading" not in name_val.lower():
            continue
        rPr = style_el.find(qn("w:rPr"))
        if rPr is None:
            continue
        color = rPr.find(qn("w:color"))
        assert color is None, f"Style '{name_val}' should not have a <w:color> element"


def test_rendered_headings_have_no_blue_color(tmp_path: Path):
    """Heading styles in rendered documents must not carry the blue theme
    color from the default python-docx template.  Both ``styles.xml`` and
    ``stylesWithEffects.xml`` are checked."""
    out = tmp_path / "heading-color.docx"
    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {
                "Heading1": {"style_id": "Heading1", "name": "Heading 1", "type": "paragraph", "based_on": "Normal"},
            },
            "body": [
                {"id": "p0", "type": "Paragraph", "style": "Heading1",
                 "content": [{"type": "Text", "text": "Title"}]},
            ],
            "passthrough": {},
        },
    }
    render_ast(ast, out)

    # Check styles.xml via python-docx API
    rebuilt = Document(out)
    _assert_no_heading_colors_in_styles(rebuilt.styles.element)

    # Check stylesWithEffects.xml via raw ZIP
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(out) as zf:
        if "word/stylesWithEffects.xml" in zf.namelist():
            swe_root = etree.fromstring(zf.read("word/stylesWithEffects.xml"))
            _assert_no_heading_colors_in_styles(swe_root)


def test_rendered_compat_mode_is_15(tmp_path: Path):
    """Rendered documents should use compatibilityMode 15 (Word 2013+)
    so that modern Word does not enter compatibility mode."""
    out = tmp_path / "compat.docx"
    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {},
            "body": [
                {"id": "p0", "type": "Paragraph", "style": None,
                 "content": [{"type": "Text", "text": "hello"}]},
            ],
            "passthrough": {},
        },
    }
    render_ast(ast, out)

    rebuilt = Document(out)
    settings_el = rebuilt.settings.element
    compat = settings_el.find(qn("w:compat"))
    assert compat is not None
    for cs in compat.iterchildren(qn("w:compatSetting")):
        if (
            cs.get(qn("w:name")) == "compatibilityMode"
            and cs.get(qn("w:uri")) == "http://schemas.microsoft.com/office/word"
        ):
            assert cs.get(qn("w:val")) == "15"
            return
    raise AssertionError("compatibilityMode setting not found")


def test_roundtrip_heading_theme_color_not_reapplied(tmp_path: Path):
    """Round-tripping a document whose heading styles use a theme color must
    NOT re-apply that color as an explicit run-level override.  The default
    python-docx template defines headings with blue theme colours; after a
    round-trip the heading runs should carry no explicit ``<w:color>`` so
    that the cleaned-up style (which has no colour) determines the text
    colour (black / auto).
    """
    src = tmp_path / "theme-heading.docx"
    out = tmp_path / "theme-heading-out.docx"

    doc = Document()
    h1 = doc.add_paragraph("Title Level 1")
    h1.style = "Heading 1"
    h2 = doc.add_paragraph("Title Level 2")
    h2.style = "Heading 2"
    doc.save(src)

    ast = parse_docx(src)

    # The parser must NOT capture the template theme colour in default_run
    for block in ast["document"]["body"]:
        if block.get("style") in ("Heading1", "Heading2"):
            assert "color" not in block.get("default_run", {}), (
                f"default_run for {block['style']} should not contain a "
                f"theme-derived color, got {block.get('default_run')}"
            )

    render_ast(ast, out)

    rebuilt = Document(out)
    for para in rebuilt.paragraphs:
        if para.style.name in ("Heading 1", "Heading 2"):
            for run in para.runs:
                rPr = run._element.find(qn("w:rPr"))
                if rPr is not None:
                    color_el = rPr.find(qn("w:color"))
                    assert color_el is None, (
                        f"Run in '{para.style.name}' should not have an "
                        f"explicit <w:color>, but found val="
                        f"{color_el.get(qn('w:val')) if color_el is not None else None}"
                    )


def test_roundtrip_table_style_preserved(tmp_path: Path):
    """Table style (e.g. 'Table Grid') must survive a parse → render round-trip
    so that cell borders remain visible in the rendered document."""
    src = tmp_path / "table_style.docx"
    out = tmp_path / "table_style_out.docx"

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.save(src)

    ast = parse_docx(src)
    table_ast = next(b for b in ast["document"]["body"] if b["type"] == "Table")
    assert table_ast.get("style") == "TableGrid", (
        f"Expected style 'TableGrid', got {table_ast.get('style')!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert rebuilt.tables[0].style.name == "Table Grid"


def test_roundtrip_inline_image_preserved(tmp_path: Path):
    """Inline images embedded in paragraphs must survive a parse → render
    round-trip; the rendered document must contain exactly one inline shape."""
    src = tmp_path / "image.docx"
    out = tmp_path / "image_out.docx"

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(io.BytesIO(_PNG_1X1), width=Inches(1), height=Inches(1))
    doc.save(src)

    ast = parse_docx(src)
    para_ast = ast["document"]["body"][0]
    image_nodes = [c for c in para_ast["content"] if c.get("type") == "InlineImage"]
    assert len(image_nodes) == 1, "Parser must produce exactly one InlineImage node"
    node = image_nodes[0]
    assert "data" in node and node["data"]
    assert node.get("width", 0) > 0
    assert node.get("height", 0) > 0

    render_ast(ast, out)

    rebuilt = Document(out)
    assert len(rebuilt.inline_shapes) == 1, (
        "Rendered document must contain exactly one inline shape"
    )


def test_table_style_resolved_by_name_when_id_differs(tmp_path: Path):
    """When the AST table style ID does not exist in the python-docx template
    (e.g. 'a8' in Chinese Word documents mapping to 'Table Grid'), the renderer
    must fall back to looking up the style by name via the styles dict."""
    out = tmp_path / "table_name.docx"

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {
                "a8": {"style_id": "a8", "name": "Table Grid", "type": "table", "based_on": None},
            },
            "body": [
                {
                    "id": "t0", "type": "Table", "style": "a8",
                    "rows": [
                        {"cells": [
                            {"id": "t0.r0c0", "content": [
                                {"id": "t0.r0c0.p0", "type": "Paragraph", "style": None,
                                 "content": [{"type": "Text", "text": "cell"}]}
                            ], "col_span": 1, "row_span": 1},
                        ]},
                    ],
                },
            ],
            "passthrough": {},
        },
    }

    render_ast(ast, out)

    rebuilt = Document(out)
    assert rebuilt.tables[0].style.name == "Table Grid"


def test_roundtrip_table_cell_font_preserved(tmp_path: Path):
    """Table cell text with explicit East Asian font must preserve the font
    through a parse → render round-trip so that CJK characters render correctly."""
    src = tmp_path / "cell_font.docx"
    out = tmp_path / "cell_font_out.docx"

    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    # Clear default paragraph and add one with styled run
    p = cell.paragraphs[0]
    run = p.add_run("发展")
    run.font.name = "宋体"
    run.font.size = Pt(12)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    runs = rebuilt.tables[0].cell(0, 0).paragraphs[0].runs
    text_runs = [r for r in runs if r.text]
    assert len(text_runs) >= 1
    r = text_runs[0]
    assert r.text == "发展"
    assert r.font.name == "宋体"
    assert _get_east_asia_font(r) == "宋体"


def test_roundtrip_preserves_paragraph_alignment(tmp_path: Path):
    """Paragraph alignment (center, right, justify) must survive a
    parse → render round-trip."""
    src = tmp_path / "align.docx"
    out = tmp_path / "align-out.docx"

    doc = Document()
    p_center = doc.add_paragraph("Centered")
    p_center.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right = doc.add_paragraph("Right")
    p_right.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_justify = doc.add_paragraph("Justified")
    p_justify.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(src)

    ast = parse_docx(src)
    body = ast["document"]["body"]
    assert body[0].get("paragraph_format", {}).get("alignment") == "center"
    assert body[1].get("paragraph_format", {}).get("alignment") == "right"
    assert body[2].get("paragraph_format", {}).get("alignment") == "justify"

    render_ast(ast, out)

    rebuilt = Document(out)
    assert rebuilt.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert rebuilt.paragraphs[1].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert rebuilt.paragraphs[2].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_roundtrip_preserves_hyperlink_text(tmp_path: Path):
    """Text inside ``<w:hyperlink>`` elements (used by TOC entries) must be
    captured by the parser so it is not lost during a round-trip."""
    src = tmp_path / "hyperlink.docx"
    out = tmp_path / "hyperlink-out.docx"

    # Build a paragraph containing a hyperlink run via raw XML manipulation
    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    hyperlink = OxmlElement("w:hyperlink")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Linked Text"
    run_el.append(t_el)
    hyperlink.append(run_el)
    p_el.append(hyperlink)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "Linked Text" in full_text, (
        f"Hyperlink text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "Linked Text" in rebuilt.paragraphs[0].text


def test_roundtrip_preserves_ins_text(tmp_path: Path):
    """Text inside ``<w:ins>`` (track-change insertion) must not be lost."""
    src = tmp_path / "ins.docx"
    out = tmp_path / "ins-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    ins = OxmlElement("w:ins")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Inserted Text"
    run_el.append(t_el)
    ins.append(run_el)
    p_el.append(ins)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "Inserted Text" in full_text, (
        f"<w:ins> text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "Inserted Text" in rebuilt.paragraphs[0].text


def test_roundtrip_preserves_fld_simple_text(tmp_path: Path):
    """Text inside ``<w:fldSimple>`` (simple field) must not be lost."""
    src = tmp_path / "fld.docx"
    out = tmp_path / "fld-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "1"
    run_el.append(t_el)
    fld.append(run_el)
    p_el.append(fld)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "1" in full_text, (
        f"<w:fldSimple> display text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "1" in rebuilt.paragraphs[0].text


def test_roundtrip_preserves_smart_tag_text(tmp_path: Path):
    """Text inside ``<w:smartTag>`` must not be lost."""
    src = tmp_path / "smart.docx"
    out = tmp_path / "smart-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    smart = OxmlElement("w:smartTag")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Smart Text"
    run_el.append(t_el)
    smart.append(run_el)
    p_el.append(smart)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "Smart Text" in full_text, (
        f"<w:smartTag> text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "Smart Text" in rebuilt.paragraphs[0].text


def test_roundtrip_preserves_inline_sdt_text(tmp_path: Path):
    """Text inside inline ``<w:sdt>`` (content control) must not be lost."""
    src = tmp_path / "sdt_inline.docx"
    out = tmp_path / "sdt_inline-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Content Control Text"
    run_el.append(t_el)
    sdt_content.append(run_el)
    sdt.append(sdt_content)
    p_el.append(sdt)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "Content Control Text" in full_text, (
        f"Inline <w:sdt> text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "Content Control Text" in rebuilt.paragraphs[0].text


def test_roundtrip_preserves_block_sdt_text(tmp_path: Path):
    """Paragraphs inside block-level ``<w:sdt>`` (content control) must not
    be lost during a round-trip."""
    src = tmp_path / "sdt_block.docx"
    out = tmp_path / "sdt_block-out.docx"

    doc = Document()
    # Add a normal paragraph first so the document isn't empty
    doc.add_paragraph("Before SDT")

    # Build a block-level <w:sdt> wrapping a paragraph
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Inside Block SDT"
    r_el.append(t_el)
    p_el.append(r_el)
    sdt_content.append(p_el)
    sdt.append(sdt_content)
    doc.element.body.append(sdt)

    doc.add_paragraph("After SDT")
    doc.save(src)

    ast = parse_docx(src)
    all_text = " ".join(
        c.get("text", "")
        for block in ast["document"]["body"]
        for c in block.get("content", [])
    )
    assert "Inside Block SDT" in all_text, (
        f"Block-level <w:sdt> text must be captured; got: {all_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    full = " ".join(p.text for p in rebuilt.paragraphs)
    assert "Inside Block SDT" in full


def test_roundtrip_preserves_custom_xml_text(tmp_path: Path):
    """Text inside ``<w:customXml>`` must not be lost."""
    src = tmp_path / "cxml.docx"
    out = tmp_path / "cxml-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p_el = p._element

    cxml = OxmlElement("w:customXml")
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Custom XML Text"
    run_el.append(t_el)
    cxml.append(run_el)
    p_el.append(cxml)
    doc.save(src)

    ast = parse_docx(src)
    content = ast["document"]["body"][0]["content"]
    full_text = "".join(c.get("text", "") for c in content)
    assert "Custom XML Text" in full_text, (
        f"<w:customXml> text must be captured; got: {full_text!r}"
    )

    render_ast(ast, out)

    rebuilt = Document(out)
    assert "Custom XML Text" in rebuilt.paragraphs[0].text


def _build_toc_sdt(title_text="目录",
                   instruction='TOC \\o "1-3" \\h \\z \\u'):
    """Helper: build a block-level ``<w:sdt>`` representing a Word TOC."""
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    docPartObj = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Table of Contents")
    docPartObj.append(gallery)
    unique = OxmlElement("w:docPartUnique")
    docPartObj.append(unique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")

    # Title paragraph
    if title_text:
        title_p = OxmlElement("w:p")
        title_r = OxmlElement("w:r")
        title_t = OxmlElement("w:t")
        title_t.text = title_text
        title_r.append(title_t)
        title_p.append(title_r)
        sdtContent.append(title_p)

    # Paragraph with field begin + instruction + separate
    field_p = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1.append(fc1)
    field_p.append(r1)

    r2 = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instruction} "
    r2.append(it)
    field_p.append(r2)

    r3 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3.append(fc2)
    field_p.append(r3)
    sdtContent.append(field_p)

    # Dummy TOC entry (generated content)
    entry_p = OxmlElement("w:p")
    entry_r = OxmlElement("w:r")
    entry_t = OxmlElement("w:t")
    entry_t.text = "Chapter 1.....1"
    entry_r.append(entry_t)
    entry_p.append(entry_r)
    sdtContent.append(entry_p)

    # Field end
    end_p = OxmlElement("w:p")
    r4 = OxmlElement("w:r")
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r4.append(fc3)
    end_p.append(r4)
    sdtContent.append(end_p)

    sdt.append(sdtContent)
    return sdt


def test_roundtrip_toc_produces_native_field(tmp_path: Path):
    """A document with a TOC SDT must round-trip into a native Word TOC
    field (fldChar begin/separate/end + instrText) wrapped in an SDT with
    ``docPartGallery = 'Table of Contents'``, not as fake formatted paragraphs."""
    src = tmp_path / "toc.docx"
    out = tmp_path / "toc-out.docx"

    doc = Document()
    doc.add_paragraph("Before TOC")

    # Insert TOC SDT before sectPr
    sdt = _build_toc_sdt()
    sectPr = doc.element.body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(sdt)
    else:
        doc.element.body.append(sdt)

    doc.add_paragraph("Chapter 1", style="Heading 1")
    doc.add_paragraph("Body text")
    doc.save(src)

    ast = parse_docx(src)

    # Parser must produce a TOC node (not plain paragraphs)
    toc_blocks = [b for b in ast["document"]["body"] if b["type"] == "TOC"]
    assert len(toc_blocks) == 1, (
        f"Expected exactly one TOC block, got {len(toc_blocks)}"
    )
    toc = toc_blocks[0]
    assert "TOC" in toc["instruction"]
    title = toc.get("title")
    assert title is not None, "TOC title must be present"
    assert title["content"][0]["text"] == "目录"

    render_ast(ast, out)

    # Verify the rendered document has a native TOC SDT with field codes
    rebuilt = Document(out)
    body = rebuilt.element.body
    sdt_els = [c for c in body if c.tag == qn("w:sdt")]
    assert len(sdt_els) == 1, "Rendered document must contain exactly one SDT"

    sdt_el = sdt_els[0]
    # Check gallery
    gallery = sdt_el.find(f".//{qn('w:docPartGallery')}")
    assert gallery is not None
    assert gallery.get(qn("w:val")) == "Table of Contents"

    # Check field codes
    fld_types = [
        fc.get(qn("w:fldCharType"))
        for fc in sdt_el.iter(qn("w:fldChar"))
    ]
    assert "begin" in fld_types, "TOC must have fldChar begin"
    assert "separate" in fld_types, "TOC must have fldChar separate"
    assert "end" in fld_types, "TOC must have fldChar end"

    # Check instruction
    instr_texts = [it.text for it in sdt_el.iter(qn("w:instrText")) if it.text]
    instr = "".join(instr_texts)
    assert "TOC" in instr, f"instrText must contain TOC, got {instr!r}"

    # Check dirty flag for auto-update
    begin_fc = next(
        fc for fc in sdt_el.iter(qn("w:fldChar"))
        if fc.get(qn("w:fldCharType")) == "begin"
    )
    assert begin_fc.get(qn("w:dirty")) == "true", (
        "TOC field begin must be marked dirty for auto-update"
    )


def test_roundtrip_toc_without_title(tmp_path: Path):
    """A TOC SDT that has no title paragraph must still round-trip correctly."""
    src = tmp_path / "toc_no_title.docx"
    out = tmp_path / "toc_no_title-out.docx"

    doc = Document()
    sdt = _build_toc_sdt(title_text=None)
    sectPr = doc.element.body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(sdt)
    else:
        doc.element.body.append(sdt)

    doc.add_paragraph("Chapter 1", style="Heading 1")
    doc.save(src)

    ast = parse_docx(src)
    toc_blocks = [b for b in ast["document"]["body"] if b["type"] == "TOC"]
    assert len(toc_blocks) == 1
    assert "title" not in toc_blocks[0]

    render_ast(ast, out)

    rebuilt = Document(out)
    body = rebuilt.element.body
    sdt_els = [c for c in body if c.tag == qn("w:sdt")]
    assert len(sdt_els) == 1
    fld_types = [
        fc.get(qn("w:fldCharType"))
        for fc in sdt_els[0].iter(qn("w:fldChar"))
    ]
    assert "begin" in fld_types
    assert "end" in fld_types


def test_roundtrip_non_toc_sdt_still_unwrapped(tmp_path: Path):
    """Block-level SDTs that are NOT a TOC must still be unwrapped into
    individual paragraphs (regression guard for existing behaviour)."""
    src = tmp_path / "sdt_plain.docx"
    out = tmp_path / "sdt_plain-out.docx"

    doc = Document()
    doc.add_paragraph("Before SDT")

    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    p_el = OxmlElement("w:p")
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "Inside Plain SDT"
    r_el.append(t_el)
    p_el.append(r_el)
    sdt_content.append(p_el)
    sdt.append(sdt_content)
    doc.element.body.append(sdt)

    doc.add_paragraph("After SDT")
    doc.save(src)

    ast = parse_docx(src)
    toc_blocks = [b for b in ast["document"]["body"] if b["type"] == "TOC"]
    assert len(toc_blocks) == 0, "Plain SDT must not be treated as TOC"

    all_text = " ".join(
        c.get("text", "")
        for block in ast["document"]["body"]
        for c in block.get("content", [])
    )
    assert "Inside Plain SDT" in all_text

    render_ast(ast, out)
    rebuilt = Document(out)
    full = " ".join(p.text for p in rebuilt.paragraphs)
    assert "Inside Plain SDT" in full


def test_roundtrip_heading_runs_no_fake_bold(tmp_path: Path):
    """Heading runs without explicit bold must NOT receive an explicit
    ``<w:b>`` element after round-trip.  The bold appearance should come
    from the heading style, not from run-level overrides (fake-bold)."""
    src = tmp_path / "fake_bold.docx"
    out = tmp_path / "fake_bold-out.docx"

    doc = Document()
    h = doc.add_paragraph("Heading Text")
    h.style = "Heading 1"
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    para = rebuilt.paragraphs[0]
    assert para.style.name == "Heading 1"
    for run in para.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            b_el = rPr.find(qn("w:b"))
            assert b_el is None, (
                "Run in heading should not have explicit <w:b>; "
                "bold must come from the paragraph style"
            )


def test_roundtrip_no_invalid_style_references(tmp_path: Path):
    """Raw pPr/rPr must NOT carry ``<w:pStyle>`` or ``<w:rStyle>``
    references that could point to non-existent styles in the target
    document, which would cause Word to report file corruption."""
    src = tmp_path / "style_ref.docx"
    out = tmp_path / "style_ref-out.docx"

    doc = Document()
    p = doc.add_paragraph("Normal text")
    r = p.add_run(" bold")
    r.bold = True
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    body = rebuilt.element.body
    for r_el in body.iter(qn("w:r")):
        rPr = r_el.find(qn("w:rPr"))
        if rPr is not None:
            rStyle = rPr.find(qn("w:rStyle"))
            assert rStyle is None, (
                "Run rPr should not contain <w:rStyle> after rendering"
            )

    for p_el in body.iter(qn("w:p")):
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            assert numPr is None, (
                "Paragraph pPr should not contain <w:numPr> after rendering"
            )


def test_roundtrip_raw_pPr_preserves_formatting(tmp_path: Path):
    """Raw pPr formatting (alignment, spacing) must be preserved even though
    ``<w:pStyle>`` is stripped from the raw XML."""
    src = tmp_path / "raw_pPr.docx"
    out = tmp_path / "raw_pPr-out.docx"

    doc = Document()
    p = doc.add_paragraph("Centered")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    doc.save(src)

    ast = parse_docx(src)
    # Ensure raw pPr is captured
    pf = ast["document"]["body"][0].get("paragraph_format", {})
    assert "_raw_pPr" in pf, "Parser must capture _raw_pPr"

    render_ast(ast, out)

    rebuilt = Document(out)
    para = rebuilt.paragraphs[0]
    assert para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert para.text == "Centered"


def test_roundtrip_explicit_bold_preserved(tmp_path: Path):
    """A run with explicit bold must preserve it through a round-trip
    via the ``_raw_rPr`` mechanism (not paragraph_defaults)."""
    src = tmp_path / "explicit_bold.docx"
    out = tmp_path / "explicit_bold-out.docx"

    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Bold Text")
    r.bold = True
    doc.save(src)

    ast = parse_docx(src)
    render_ast(ast, out)

    rebuilt = Document(out)
    run = rebuilt.paragraphs[0].runs[0]
    assert run.text == "Bold Text"
    assert run.bold is True


def test_rendered_pPr_has_no_sectPr(tmp_path: Path):
    """Raw pPr must NOT carry ``<w:sectPr>`` into the rendered document.

    Multi-section source documents store section properties inside the last
    paragraph's ``<w:pPr>``.  If this sectPr is injected into the rendered
    document, its header/footer relationship IDs point to wrong targets,
    causing Word to report *unreadable content* and trigger recovery, which
    strips the entire pPr and loses paragraph formatting (alignment, spacing).
    """
    out = tmp_path / "sect-out.docx"

    # Build an AST whose raw pPr contains a sectPr with a bogus
    # header reference — exactly what the parser produces for the
    # last paragraph of a section in a multi-section source document.
    raw_pPr = (
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:jc w:val="center"/>'
        '<w:sectPr>'
        '<w:headerReference w:type="default" r:id="rId99"/>'
        '<w:pgSz w:w="12240" w:h="15840"/>'
        '</w:sectPr>'
        '</w:pPr>'
    )

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {},
            "body": [
                {
                    "id": "p0", "type": "Paragraph", "style": None,
                    "paragraph_format": {
                        "alignment": "center",
                        "_raw_pPr": raw_pPr,
                    },
                    "content": [{"type": "Text", "text": "Section 1"}],
                },
                {
                    "id": "p1", "type": "Paragraph", "style": None,
                    "content": [{"type": "Text", "text": "Section 2"}],
                },
            ],
            "passthrough": {},
        },
    }

    render_ast(ast, out)

    rebuilt = Document(out)
    body = rebuilt.element.body
    # No pPr in the rendered document should contain a sectPr
    for p_el in body.iter(qn("w:p")):
        p_pPr = p_el.find(qn("w:pPr"))
        if p_pPr is not None:
            assert p_pPr.find(qn("w:sectPr")) is None, (
                "Rendered pPr must not contain <w:sectPr>"
            )

    # Alignment must survive (not lost due to recovery)
    assert rebuilt.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_rendered_pPr_has_no_pPrChange(tmp_path: Path):
    """Raw pPr must NOT carry ``<w:pPrChange>`` (track-changes data) into
    the rendered document, as it may reference non-existent revision IDs."""
    out = tmp_path / "pPrChange.docx"

    raw_pPr = (
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:jc w:val="center"/>'
        '<w:pPrChange w:id="1" w:author="test" w:date="2025-01-01T00:00:00Z">'
        '<w:pPr><w:jc w:val="left"/></w:pPr>'
        '</w:pPrChange>'
        '</w:pPr>'
    )

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {},
            "body": [
                {
                    "id": "p0", "type": "Paragraph", "style": None,
                    "paragraph_format": {
                        "alignment": "center",
                        "_raw_pPr": raw_pPr,
                    },
                    "content": [{"type": "Text", "text": "Center"}],
                },
            ],
            "passthrough": {},
        },
    }

    render_ast(ast, out)

    rebuilt = Document(out)
    body = rebuilt.element.body
    for p_el in body.iter(qn("w:p")):
        p_pPr = p_el.find(qn("w:pPr"))
        if p_pPr is not None:
            assert p_pPr.find(qn("w:pPrChange")) is None, (
                "Rendered pPr must not contain <w:pPrChange>"
            )

    # Alignment must be preserved
    assert rebuilt.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_rendered_rPr_has_no_rPrChange(tmp_path: Path):
    """Raw rPr must NOT carry ``<w:rPrChange>`` (track-changes data) into
    the rendered document."""
    out = tmp_path / "rPrChange.docx"

    raw_rPr = (
        '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:b/>'
        '<w:rPrChange w:id="2" w:author="test" w:date="2025-01-01T00:00:00Z">'
        '<w:rPr/>'
        '</w:rPrChange>'
        '</w:rPr>'
    )

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": {},
            "styles": {},
            "body": [
                {
                    "id": "p0", "type": "Paragraph", "style": None,
                    "content": [
                        {"type": "Text", "text": "Bold",
                         "overrides": {"bold": True, "_raw_rPr": raw_rPr}},
                    ],
                },
            ],
            "passthrough": {},
        },
    }

    render_ast(ast, out)

    rebuilt = Document(out)
    body = rebuilt.element.body
    for r_el in body.iter(qn("w:r")):
        rPr = r_el.find(qn("w:rPr"))
        if rPr is not None:
            assert rPr.find(qn("w:rPrChange")) is None, (
                "Rendered rPr must not contain <w:rPrChange>"
            )

    assert rebuilt.paragraphs[0].runs[0].bold is True
