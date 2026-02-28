import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .paragraph_parser import parse_paragraph_block
from .style_parser import parse_styles
from .table_parser import parse_table_block


def _parse_meta(doc) -> dict:
    section = doc.sections[0]
    return {
        "page": {
            "size": "custom",
            "width": section.page_width.twips,
            "height": section.page_height.twips,
            "orientation": "landscape" if section.page_width > section.page_height else "portrait",
            "margin": {
                "top": section.top_margin.twips,
                "bottom": section.bottom_margin.twips,
                "left": section.left_margin.twips,
                "right": section.right_margin.twips,
            },
        },
        "default_style": "Normal",
        "language": "zh-CN",
    }


def parse_docx(input_path: str | Path, output_dir: str | Path | None = None) -> dict:
    doc = Document(str(input_path))
    body = []
    p_i = 0
    t_i = 0

    _tag_sdt = qn("w:sdt")
    _tag_sdt_content = qn("w:sdtContent")

    def _process_body_element(child):
        nonlocal p_i, t_i
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None:
                paragraph = Paragraph(child, doc)
            body.append(parse_paragraph_block(paragraph, f"p{p_i}"))
            p_i += 1
        elif tag == "tbl":
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is None:
                table = Table(child, doc)
            body.append(parse_table_block(table, f"t{t_i}"))
            t_i += 1

    for child in doc.element.body:
        if child.tag == _tag_sdt:
            sdt_content = child.find(_tag_sdt_content)
            if sdt_content is not None:
                for inner in sdt_content:
                    _process_body_element(inner)
        elif child.tag == qn("w:sectPr"):
            continue
        else:
            _process_body_element(child)

    ast = {
        "schema_version": "1.0",
        "document": {
            "meta": _parse_meta(doc),
            "styles": parse_styles(doc),
            "body": body,
            "passthrough": {},
        },
    }

    if output_dir:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        (out_dir / "document.ast.json").write_text(json.dumps(ast, ensure_ascii=False, indent=2), encoding="utf-8")
        (out_dir / "media").mkdir(exist_ok=True)

    return ast
